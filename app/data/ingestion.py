import os
from typing import List
from app.data.chunking import chunk_text
from app.models.embeddings import EmbeddingClient
from app.utils import config
from app.utils.logger import logger

from langchain_chroma import Chroma
from langchain.schema import Document

from PyPDF2 import PdfReader
import docx


from app.data.marker import extract_text_with_marker

def read_pdf(path: str) -> str:
    reader = PdfReader(path)
    text = []
    for page in reader.pages:
        try:
            t = page.extract_text()
            text.append(t or "")
        except Exception:
            text.append("")
    return "\n".join(text)

def read_docx(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def load_file_to_text(path: str, use_marker_ocr: bool = True) -> str:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        # Primero intentamos extraer texto normal
        raw_text = read_pdf(path)
        # Si el texto es muy corto o parece vacío, o si forzamos OCR:
        if use_marker_ocr:
            # criterio simple: digamos si menos de N caracteres
            if len(raw_text.strip()) < 500:  
                logger.info(f"PDF {path} parece tener poco texto (" \
                            f"{len(raw_text)} chars). Usando OCR de Marker.")
                try:
                    return extract_text_with_marker(path, force_ocr=True)
                except Exception as e:
                    logger.exception(f"Marker OCR falló en {path}: {e}")
                    # fallback al texto crudo
                    return raw_text
        return raw_text

    elif ext in [".docx", ".doc"]:
        return read_docx(path)
    elif ext in [".txt", ".md"]:
        return read_txt(path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")

def ingest_files(paths: List[str], collection_name: str = "study_collection", persist: bool = True):
    def clean_text(text):
        return text.encode('utf-8', 'ignore').decode('utf-8')
    logger.info("Starting ingestion pipeline with OCR-marker support...")
    emb = EmbeddingClient()
    vectordb = Chroma(persist_directory=config.CHROMA_PERSIST_DIR, embedding_function=emb._client)
    documents = []
    for path in paths:
        try:
            text = load_file_to_text(path, use_marker_ocr=True)
            if not text or len(text.strip()) == 0:
                logger.warning(f"No se extrajo texto de {path}")
                continue
            chunks = chunk_text(text, chunk_size_chars=config.MAX_CHUNK_SIZE)
            for i, ch in enumerate(chunks):
                ch_clean = clean_text(ch)
                metadata = {"source": os.path.basename(path), "chunk": i}
                documents.append(Document(page_content=ch_clean, metadata=metadata))
            logger.info(f"Ingested {len(chunks)} chunks from {path}")
        except Exception as e:
            logger.exception(f"Failed to ingest {path}: {e}")

    if documents:
        vectordb.add_documents(documents)
        logger.info(f"Added {len(documents)} documents to Chroma at {config.CHROMA_PERSIST_DIR}")
    else:
        logger.warning("No documents were created from ingestion.")
